from __future__ import annotations

import io
import sqlite3
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from langchain_core.documents import Document
from PIL import Image, ImageDraw, ImageFont

import build_index
from document_service import (
    DocumentService,
    DocumentValidationError,
    normalize_text,
    sanitize_filename,
)


def _docx_bytes() -> bytes:
    from docx import Document

    document = Document()
    document.add_heading("液压泵资料", level=1)
    document.add_paragraph("适用机型 PC200-8。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "零件号"
    table.cell(0, 1).text = "型号"
    table.cell(1, 0).text = "708-2L-00500"
    table.cell(1, 1).text = "PC200-8"
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _xlsx_bytes() -> bytes:
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "库存"
    sheet.append(["零件号", "数量"])
    sheet.append(["708-2L-00500", 3])
    output = io.BytesIO()
    workbook.save(output)
    workbook.close()
    return output.getvalue()


def _pdf_bytes(text: str = "Hydraulic pump PC200") -> bytes:
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 16 Tf 72 720 Td ({escaped}) Tj ET".encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>"
        ),
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n"
        + stream
        + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for number, obj in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{number} 0 obj\n".encode("ascii"))
        output.extend(obj)
        output.extend(b"\nendobj\n")
    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(output)


def _image_bytes(image_format: str, text: str = "PC200 HYDRAULIC PUMP") -> bytes:
    image = Image.new("RGB", (1100, 220), "white")
    draw = ImageDraw.Draw(image)
    font_path = Path("C:/Windows/Fonts/arial.ttf")
    font = ImageFont.truetype(str(font_path), 72) if font_path.exists() else None
    draw.text((35, 60), text, fill="black", font=font)
    output = io.BytesIO()
    image.save(output, format=image_format, quality=95)
    return output.getvalue()


class DocumentServiceTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary_directory = tempfile.TemporaryDirectory()
        self.root = Path(self.temporary_directory.name)
        self.service = DocumentService(
            db_path=self.root / "rag.sqlite3",
            upload_dir=self.root / "uploads",
            document_max_bytes=1024 * 1024,
            image_max_bytes=3 * 1024 * 1024,
            batch_max_bytes=5 * 1024 * 1024,
        )

    def tearDown(self) -> None:
        self.temporary_directory.cleanup()

    def test_utf8_txt_parsing_and_normalization(self) -> None:
        data = "标题\r\n\r\n\r\nPC200-8 液压泵 12.5-A".encode("utf-8")
        parsed = self.service.parse_content(".txt", data)
        self.assertEqual("标题\n\nPC200-8 液压泵 12.5-A", parsed.normalized_text)
        self.assertEqual("plain_text", parsed.parser_name)

    def test_non_utf8_chinese_txt_parsing(self) -> None:
        parsed = self.service.parse_content(
            ".txt",
            "液压泵适用机型 PC200-8".encode("gb18030"),
        )
        self.assertIn("液压泵适用机型", parsed.normalized_text)
        self.assertEqual("gb18030", parsed.segments[0]["encoding"])

    def test_docx_paragraph_and_heading_parsing(self) -> None:
        parsed = self.service.parse_content(".docx", _docx_bytes())
        self.assertIn("# 液压泵资料", parsed.normalized_text)
        self.assertIn("适用机型 PC200-8", parsed.normalized_text)

    def test_docx_table_parsing(self) -> None:
        parsed = self.service.parse_content(".docx", _docx_bytes())
        table_segment = next(
            item for item in parsed.segments if item.get("section_type") == "table"
        )
        self.assertIn("C1=零件号", table_segment["text"])
        self.assertIn("C2=PC200-8", table_segment["text"])
        self.assertEqual((1, 2), (table_segment["row_start"], table_segment["row_end"]))

    def test_xlsx_preserves_sheet_rows_and_coordinates(self) -> None:
        parsed = self.service.parse_content(".xlsx", _xlsx_bytes())
        self.assertIn("## Sheet: 库存", parsed.normalized_text)
        self.assertIn("Row 2: A2=708-2L-00500 | B2=3", parsed.normalized_text)
        self.assertEqual("库存", parsed.segments[0]["sheet"])
        self.assertEqual((1, 2), (
            parsed.segments[0]["row_start"],
            parsed.segments[0]["row_end"],
        ))

    def test_pdf_parser_returns_nonempty_page_segment(self) -> None:
        parsed = self.service.parse_content(".pdf", _pdf_bytes())
        self.assertIn("Hydraulic pump PC200", parsed.normalized_text)
        self.assertEqual(1, parsed.segments[0]["page"])

    def test_png_and_jpg_run_real_local_ocr(self) -> None:
        for extension, image_format in ((".png", "PNG"), (".jpg", "JPEG")):
            with self.subTest(extension=extension):
                parsed = self.service.parse_content(
                    extension,
                    _image_bytes(image_format),
                )
                compact = parsed.normalized_text.upper().replace(" ", "")
                self.assertIn("PC200", compact)
                self.assertIn("PUMP", compact)
                self.assertEqual("rapidocr_onnxruntime", parsed.parser_name)

    def test_oversized_file_is_rejected(self) -> None:
        tiny_limit_service = DocumentService(
            db_path=self.root / "small.sqlite3",
            upload_dir=self.root / "small_uploads",
            document_max_bytes=4,
        )
        with self.assertRaises(DocumentValidationError):
            tiny_limit_service.upload_file("large.txt", b"12345")

    def test_disallowed_extension_is_rejected(self) -> None:
        with self.assertRaises(DocumentValidationError):
            self.service.upload_file("script.exe", b"MZ")

    def test_path_traversal_filename_is_rejected(self) -> None:
        for filename in ("../secret.txt", r"..\secret.txt", "C:secret.txt"):
            with self.subTest(filename=filename):
                with self.assertRaises(DocumentValidationError):
                    sanitize_filename(filename)

    def test_identical_file_reuses_existing_version_without_reparse(self) -> None:
        data = "PC200 液压泵".encode("utf-8")
        with patch.object(
            self.service,
            "parse_content",
            wraps=self.service.parse_content,
        ) as parser:
            first = self.service.upload_file("pump.txt", data)
            second = self.service.upload_file("pump.txt", data)
        self.assertEqual("uploaded", first.status)
        self.assertEqual("duplicate_file", second.status)
        self.assertEqual(first.version_id, second.version_id)
        self.assertEqual(1, parser.call_count)
        self.assertEqual(1, len(self.service.list_versions()))

    def test_same_name_different_content_creates_new_active_version(self) -> None:
        first = self.service.upload_file("pump.txt", "版本一".encode("utf-8"))
        second = self.service.upload_file("pump.txt", "版本二".encode("utf-8"))
        versions = self.service.list_versions(first.document_id)
        by_number = {item["version_number"]: item for item in versions}
        self.assertEqual("new_version", second.status)
        self.assertEqual(2, second.version_number)
        self.assertEqual(first.document_id, second.document_id)
        self.assertEqual(0, by_number[1]["is_active"])
        self.assertEqual("inactive", by_number[1]["index_status"])
        self.assertEqual(1, by_number[2]["is_active"])

    def test_different_names_same_normalized_content_is_duplicate_content(self) -> None:
        first = self.service.upload_file(
            "pump-a.txt",
            "液压泵\r\n\r\n说明".encode("utf-8"),
        )
        second = self.service.upload_file(
            "pump-b.txt",
            b"\xef\xbb\xbf" + "液压泵\n\n\n说明".encode("utf-8"),
        )
        self.assertEqual("duplicate_content", second.status)
        self.assertEqual(first.document_id, second.document_id)
        self.assertEqual(1, len(self.service.list_versions()))

    def test_parse_failure_is_recorded_but_not_loaded_as_active(self) -> None:
        blank_image = Image.new("RGB", (200, 100), "white")
        output = io.BytesIO()
        blank_image.save(output, format="PNG")
        result = self.service.upload_file("blank.png", output.getvalue())
        versions = self.service.list_versions(result.document_id)
        self.assertEqual("parse_failed", result.status)
        self.assertEqual("parse_failed", versions[0]["parse_status"])
        self.assertTrue(versions[0]["parse_error"])
        self.assertEqual([], self.service.load_active_documents())

    def test_deactivated_document_is_not_loaded_by_build_index(self) -> None:
        result = self.service.upload_file(
            "pump.txt",
            "PC200 液压泵资料".encode("utf-8"),
        )
        docs_dir = self.root / "empty_docs"
        docs_dir.mkdir()
        with patch.object(build_index, "DOCS_DIR", docs_dir):
            before = build_index.load_documents(self.service)
            self.assertTrue(any(doc.metadata.get("version_id") for doc in before))
            self.assertTrue(self.service.deactivate_document(result.document_id))
            after = build_index.load_documents(self.service)
        self.assertEqual([], after)

    def test_uploaded_document_metadata_and_index_status_updates(self) -> None:
        result = self.service.upload_file(
            "pump.txt",
            "PC200 液压泵资料".encode("utf-8"),
            uploaded_by="tester",
        )
        document = self.service.load_active_documents()[0]
        for key in (
            "document_id",
            "version_id",
            "version_number",
            "source",
            "original_name",
            "parser_name",
        ):
            self.assertIn(key, document.metadata)

        self.service.mark_indexed({result.version_id: 2})
        indexed = self.service.list_versions(result.document_id)[0]
        self.assertEqual("indexed", indexed["index_status"])
        self.assertEqual(2, indexed["chunk_count"])
        self.assertTrue(indexed["indexed_at"])

        self.service.mark_index_failed([result.version_id], RuntimeError("secret path"))
        failed = self.service.list_versions(result.document_id)[0]
        self.assertEqual("index_failed", failed["index_status"])
        self.assertTrue(failed["index_error"])

    def test_database_contains_required_tables_and_columns(self) -> None:
        connection = sqlite3.connect(self.service.db_path)
        try:
            tables = {
                row[0]
                for row in connection.execute(
                    "SELECT name FROM sqlite_master WHERE type = ?",
                    ("table",),
                )
            }
            version_columns = {
                row[1]
                for row in connection.execute(
                    "PRAGMA table_info(rag_document_versions)"
                )
            }
        finally:
            connection.close()
        self.assertIn("rag_documents", tables)
        self.assertIn("rag_document_versions", tables)
        self.assertTrue(
            {
                "file_sha256",
                "normalized_content_sha256",
                "parse_status",
                "index_status",
                "normalized_text",
                "is_active",
            }.issubset(version_columns)
        )

    def test_normalize_text_keeps_business_punctuation_and_boundaries(self) -> None:
        value = normalize_text("PC200-8\t708-2L-00500\r\n\r\n12.5 液压泵")
        self.assertEqual("PC200-8\t708-2L-00500\n\n12.5 液压泵", value)

    def test_uploaded_chunk_identity_adds_version_without_changing_static_payload(self) -> None:
        static = Document(
            page_content="PC200",
            metadata={"source": "docs/pump.md", "page": 1},
        )
        uploaded = Document(
            page_content="PC200",
            metadata={
                "source": "upload/doc/version/pump.md",
                "document_id": "doc",
                "version_id": "version",
                "version_number": 2,
                "page": 1,
            },
        )
        static_payload = build_index._chunk_identity_payload(static, 0, 500, 80)
        uploaded_payload = build_index._chunk_identity_payload(uploaded, 0, 500, 80)
        self.assertNotIn("version_id", static_payload)
        self.assertNotIn("page", static_payload)
        self.assertEqual("version", uploaded_payload["version_id"])
        self.assertEqual(1, uploaded_payload["page"])

    def test_incremental_index_marks_upload_and_removes_it_after_deactivation(self) -> None:
        class FakeVectorStore:
            def __init__(self) -> None:
                self.ids: set[str] = set()
                self.uploaded_ids: set[str] = set()

            def add_documents(self, documents, ids) -> None:
                self.ids.update(ids)
                self.uploaded_ids.update(
                    chunk_id
                    for document, chunk_id in zip(documents, ids)
                    if document.metadata.get("version_id")
                )

            def delete(self, ids) -> None:
                self.ids.difference_update(ids)

        docs_dir = self.root / "docs"
        docs_dir.mkdir()
        (docs_dir / "base.txt").write_text("基础知识", encoding="utf-8")
        chroma_dir = self.root / "chroma"
        manifest_path = chroma_dir / "index_manifest.json"
        store = FakeVectorStore()
        result = self.service.upload_file(
            "pump.txt",
            "PC200 液压泵资料".encode("utf-8"),
        )
        fingerprint = {"provider": "fake", "dimensions": 3}

        with (
            patch.object(build_index, "BASE_DIR", self.root),
            patch.object(build_index, "DOCS_DIR", docs_dir),
            patch.object(build_index, "CHROMA_DB_DIR", chroma_dir),
            patch.object(build_index, "MANIFEST_PATH", manifest_path),
            patch.object(build_index, "_get_embeddings", return_value=object()),
            patch.object(
                build_index,
                "build_index_fingerprint",
                return_value=fingerprint,
            ),
            patch.object(build_index, "create_vector_store", return_value=store),
        ):
            build_index.build_index(incremental=True, document_service=self.service)
            indexed = self.service.list_versions(result.document_id)[0]
            self.assertEqual("indexed", indexed["index_status"])
            self.assertGreater(indexed["chunk_count"], 0)
            uploaded_ids = set(store.uploaded_ids)
            self.assertTrue(uploaded_ids.issubset(store.ids))
            self.assertTrue(self.service.deactivate_document(result.document_id))
            build_index.build_index(incremental=True, document_service=self.service)

        inactive = self.service.list_versions(result.document_id)[0]
        self.assertEqual("inactive", inactive["index_status"])
        self.assertEqual(0, inactive["chunk_count"])
        self.assertTrue(uploaded_ids.isdisjoint(store.ids))


if __name__ == "__main__":
    unittest.main()
